"""ATS-friendly DOCX rendering for factual source and optimized resumes."""

import hashlib
from collections import OrderedDict
from collections.abc import Iterable, Sequence
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ai_internship_assistant.domain.models import (
    Certification,
    DocxRenderOptions,
    Education,
    Experience,
    OptimizedResume,
    OptimizedResumeContact,
    Project,
    RenderedResume,
    RenderedResumeFile,
    Resume,
    ResumeOutputFormat,
    ResumeRenderOptions,
    Skill,
)
from ai_internship_assistant.services.resume_generator import (
    MarkdownResumeRenderer,
    ResumeOutputFileExistsError,
    ResumeOutputWriteError,
    ResumeRenderingError,
    UnsupportedResumeFormatError,
)
from ai_internship_assistant.utils import canonical_skill_name
from ai_internship_assistant.utils.filenames import generate_resume_filename, sanitize_filename

_RENDERER_VERSION = "docx-resume-renderer-v1"
_DEFAULT_SECTION_ORDER = [
    "Summary",
    "Education",
    "Certifications",
    "Skills",
    "Projects",
    "Experience",
    "Additional Sections",
]


class DocxResumeRenderer:
    """Render existing resume data into a compact, table-free Word document."""

    def render(
        self,
        resume: Resume | OptimizedResume,
        options: ResumeRenderOptions | None = None,
    ) -> RenderedResume:
        """Build an in-memory DOCX and return deterministic visible-text metadata."""

        self._validate_resume(resume)
        resolved = self._options(options)
        document, warnings = self._build_document(resume, resolved)
        content = "\n".join(
            paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
        )
        metadata = MarkdownResumeRenderer().render(resume, resolved)
        return RenderedResume(
            content=content.strip() + "\n",
            format=ResumeOutputFormat.DOCX,
            candidate_name=metadata.candidate_name,
            target_job_title=metadata.target_job_title,
            target_company=metadata.target_company,
            source_resume_id=resolved.source_resume_id,
            source_version_id=resolved.source_version_id,
            warnings=self._deduplicate([*metadata.warnings, *warnings]),
            renderer_version=_RENDERER_VERSION,
        )

    def render_to_file(
        self,
        resume: Resume | OptimizedResume,
        output_path: Path,
        options: ResumeRenderOptions | None = None,
        *,
        overwrite: bool | None = None,
    ) -> RenderedResumeFile:
        """Write one ATS-friendly DOCX and return artifact metadata."""

        self._validate_resume(resume)
        resolved = self._options(options)
        rendered = self.render(resume, resolved)
        requested_directory = output_path.exists() and output_path.is_dir()
        path = self._output_path(output_path, rendered)
        allow_overwrite = resolved.allow_overwrite if overwrite is None else overwrite
        if path.exists() and not allow_overwrite:
            if requested_directory:
                path = self._next_available_path(path)
            else:
                raise ResumeOutputFileExistsError("resume output file already exists")
        document, _ = self._build_document(resume, resolved)
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            document.save(str(path))
            content = path.read_bytes()
        except (OSError, ValueError) as exc:
            raise ResumeOutputWriteError("DOCX resume output file could not be written") from exc
        return RenderedResumeFile(
            path=str(path),
            format=ResumeOutputFormat.DOCX,
            byte_size=len(content),
            content_hash=hashlib.sha256(content).hexdigest(),
            rendered_resume=rendered,
        )

    def _build_document(
        self,
        resume: Resume | OptimizedResume,
        options: DocxRenderOptions,
    ) -> tuple[DocumentObject, list[str]]:
        document = Document()
        self._configure_document(document, options)
        view = self._view(resume)
        warnings: list[str] = []
        name = view["name"]
        if isinstance(name, str) and name:
            self._add_name(document, name, options)
        else:
            warnings.append("candidate name is missing")
        contact = view["contact"]
        if options.include_contact and isinstance(contact, OptimizedResumeContact):
            if not self._add_contact(document, contact, options):
                warnings.append("contact information is missing")
        rendered_sections = 0
        for section in self._section_order(resume, options, warnings):
            rendered_sections += self._render_section(
                document,
                section,
                view,
                options,
                warnings,
            )
        if rendered_sections == 0:
            warnings.append("resume contains no renderable sections")
        if len(document.paragraphs) > 55:
            warnings.append("DOCX content may exceed one page; review the generated document")
        return document, self._deduplicate(warnings)

    def _configure_document(self, document: DocumentObject, options: DocxRenderOptions) -> None:
        section = document.sections[0]
        section.start_type = WD_SECTION.NEW_PAGE
        section.top_margin = Inches(options.margin_top)
        section.bottom_margin = Inches(options.margin_bottom)
        section.left_margin = Inches(options.margin_left)
        section.right_margin = Inches(options.margin_right)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

        normal = document.styles["Normal"]
        normal.font.name = options.font_name
        normal.font.size = Pt(options.font_size)
        normal.font.color.rgb = RGBColor(0, 0, 0)
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), options.font_name)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(2 if options.compact_spacing else 6)
        normal.paragraph_format.line_spacing = options.line_spacing
        document.core_properties.author = ""
        document.core_properties.last_modified_by = ""
        document.core_properties.title = ""

        bullet = document.styles["List Bullet"]
        bullet.font.name = options.font_name
        bullet.font.size = Pt(options.font_size)
        bullet.font.color.rgb = RGBColor(0, 0, 0)
        bullet.element.rPr.rFonts.set(qn("w:eastAsia"), options.font_name)
        bullet.paragraph_format.left_indent = Inches(options.bullet_indent)
        bullet.paragraph_format.first_line_indent = Inches(-0.188)
        bullet.paragraph_format.space_after = Pt(2 if options.compact_spacing else 4)
        bullet.paragraph_format.line_spacing = options.line_spacing

    def _add_name(
        self,
        document: DocumentObject,
        name: str,
        options: DocxRenderOptions,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(name)
        run.bold = True
        run.font.name = options.font_name
        run.font.size = Pt(options.name_font_size)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_contact(
        self,
        document: DocumentObject,
        contact: OptimizedResumeContact,
        options: DocxRenderOptions,
    ) -> bool:
        values = [
            value
            for value in [
                contact.email,
                contact.phone,
                contact.linkedin_url,
                contact.github_url,
                *contact.links,
                contact.location,
            ]
            if value
        ]
        if not values:
            return False
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4 if options.compact_spacing else 8)
        paragraph.add_run(" | ".join(values))
        return True

    def _render_section(
        self,
        document: DocumentObject,
        section: str,
        view: dict[str, object],
        options: DocxRenderOptions,
        warnings: list[str],
    ) -> int:
        name = section.casefold()
        if name == "summary":
            return self._summary(document, view["summary"], options)
        if name == "education":
            return self._education(document, view["education"], options)
        if name == "certifications":
            return self._certifications(document, view["certifications"], options)
        if name == "skills":
            return self._skills(document, view["skills"], options)
        if name == "projects":
            return self._projects(document, view["projects"], options, warnings)
        if name == "experience":
            return self._experience(document, view["experience"], options, warnings)
        if name in {"additional", "additional sections"}:
            return self._additional(document, view["additional"], options)
        return 0

    def _summary(
        self,
        document: DocumentObject,
        value: object,
        options: DocxRenderOptions,
    ) -> int:
        if not options.include_summary or not isinstance(value, str) or not value.strip():
            return 0
        self._add_section_heading(document, "SUMMARY", options)
        document.add_paragraph(value.strip())
        return 1

    def _education(
        self,
        document: DocumentObject,
        value: object,
        options: DocxRenderOptions,
    ) -> int:
        entries = (
            [item for item in value if isinstance(item, Education)]
            if isinstance(value, list)
            else []
        )
        if not options.include_education or not entries:
            return 0
        self._add_section_heading(document, "EDUCATION", options)
        for item in entries:
            qualification = ", ".join(part for part in [item.degree, item.program] if part)
            self._add_item_heading(
                document,
                " - ".join(part for part in [qualification, item.institution] if part),
                options,
            )
            dates = self._dates(item.start_date, item.end_date)
            if dates:
                document.add_paragraph(dates)
            for detail in item.details:
                document.add_paragraph(detail, style="List Bullet")
        return 1

    def _certifications(
        self,
        document: DocumentObject,
        value: object,
        options: DocxRenderOptions,
    ) -> int:
        entries = (
            [item for item in value if isinstance(item, Certification)]
            if isinstance(value, list)
            else []
        )
        if not options.include_certifications or not entries:
            return 0
        self._add_section_heading(document, "CERTIFICATIONS", options)
        for item in entries:
            text = " - ".join(part for part in [item.name, item.issuer] if part)
            dates = self._dates(item.issued_date, item.expiration_date)
            document.add_paragraph(f"{text} ({dates})" if dates else text, style="List Bullet")
        return 1

    def _skills(
        self,
        document: DocumentObject,
        value: object,
        options: DocxRenderOptions,
    ) -> int:
        skills = self._deduplicate_skills(
            item for item in value if isinstance(item, Skill)
        ) if isinstance(value, list) else []
        if not options.include_skills or not skills:
            return 0
        self._add_section_heading(document, "SKILLS", options)
        if any(skill.category for skill in skills):
            groups: OrderedDict[str, list[str]] = OrderedDict()
            for skill in skills:
                groups.setdefault(skill.category or "Other", []).append(skill.name)
            for category, names in groups.items():
                self._add_labeled_paragraph(document, category, ", ".join(names))
        else:
            document.add_paragraph(", ".join(skill.name for skill in skills))
        return 1

    def _projects(
        self,
        document: DocumentObject,
        value: object,
        options: DocxRenderOptions,
        warnings: list[str],
    ) -> int:
        entries = (
            [item for item in value if isinstance(item, Project)]
            if isinstance(value, list)
            else []
        )
        if not options.include_projects or not entries:
            return 0
        entries = self._limited_entries(entries, options.max_projects, "project", warnings)
        if not entries:
            return 0
        self._add_section_heading(document, "PROJECTS", options)
        for item in entries:
            self._add_item_heading(document, item.name, options)
            if item.description:
                document.add_paragraph(item.description)
            if item.technologies:
                self._add_labeled_paragraph(document, "Technologies", ", ".join(item.technologies))
            for bullet in self._limited_bullets(
                item.bullets,
                options.max_bullets_per_project,
                f"project {item.name}",
                warnings,
            ):
                document.add_paragraph(bullet, style="List Bullet")
            if item.links:
                self._add_labeled_paragraph(document, "Links", ", ".join(item.links))
        return 1

    def _experience(
        self,
        document: DocumentObject,
        value: object,
        options: DocxRenderOptions,
        warnings: list[str],
    ) -> int:
        entries = (
            [item for item in value if isinstance(item, Experience)]
            if isinstance(value, list)
            else []
        )
        if not options.include_experience or not entries:
            return 0
        entries = self._limited_entries(entries, options.max_experiences, "experience", warnings)
        if not entries:
            return 0
        self._add_section_heading(document, "EXPERIENCE", options)
        for item in entries:
            self._add_item_heading(document, f"{item.title} - {item.organization}", options)
            details = " | ".join(
                part
                for part in [item.location, self._dates(item.start_date, item.end_date)]
                if part
            )
            if details:
                document.add_paragraph(details)
            for bullet in self._limited_bullets(
                item.bullets,
                options.max_bullets_per_experience,
                f"experience {item.title} at {item.organization}",
                warnings,
            ):
                document.add_paragraph(bullet, style="List Bullet")
            if item.technologies:
                self._add_labeled_paragraph(document, "Technologies", ", ".join(item.technologies))
        return 1

    def _additional(
        self,
        document: DocumentObject,
        value: object,
        options: DocxRenderOptions,
    ) -> int:
        if not options.include_additional_sections or not isinstance(value, dict):
            return 0
        rendered = 0
        for title, entries in value.items():
            if isinstance(title, str) and isinstance(entries, list) and entries:
                self._add_section_heading(document, title.upper(), options)
                for entry in entries:
                    document.add_paragraph(str(entry), style="List Bullet")
                rendered += 1
        return rendered

    def _add_section_heading(
        self,
        document: DocumentObject,
        text: str,
        options: DocxRenderOptions,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(6 if options.compact_spacing else 10)
        paragraph.paragraph_format.space_after = Pt(2 if options.compact_spacing else 4)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = options.font_name
        run.font.size = Pt(options.heading_font_size)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_item_heading(
        self,
        document: DocumentObject,
        text: str,
        options: DocxRenderOptions,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = options.font_name
        run.font.size = Pt(options.font_size)

    def _add_labeled_paragraph(
        self,
        document: DocumentObject,
        label: str,
        value: str,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    def _view(self, resume: Resume | OptimizedResume) -> dict[str, object]:
        if isinstance(resume, OptimizedResume):
            return {
                "name": resume.contact.full_name,
                "contact": resume.contact,
                "summary": resume.summary,
                "education": resume.education,
                "certifications": resume.certifications,
                "skills": resume.skills,
                "projects": resume.projects,
                "experience": resume.experience,
                "additional": resume.additional_sections,
            }
        return {
            "name": resume.full_name,
            "contact": OptimizedResumeContact(
                source_file=resume.source_file,
                full_name=resume.full_name,
                email=resume.email,
                phone=resume.phone,
                location=resume.location,
                linkedin_url=resume.linkedin_url,
                github_url=resume.github_url,
                links=list(resume.links),
            ),
            "summary": resume.summary,
            "education": resume.education,
            "certifications": resume.certifications,
            "skills": resume.skills,
            "projects": resume.projects,
            "experience": resume.experience,
            "additional": {},
        }

    def _section_order(
        self,
        resume: Resume | OptimizedResume,
        options: DocxRenderOptions,
        warnings: list[str],
    ) -> list[str]:
        source = options.section_order
        if source is None and isinstance(resume, OptimizedResume):
            source = resume.metadata.section_order
        order = list(source or _DEFAULT_SECTION_ORDER)
        known = {item.casefold() for item in _DEFAULT_SECTION_ORDER} | {"contact", "additional"}
        unknown = [item for item in order if item.casefold() not in known]
        if unknown:
            warnings.append(f"ignored unknown sections: {', '.join(unknown)}")
        included = {item.casefold() for item in order}
        order.extend(item for item in _DEFAULT_SECTION_ORDER if item.casefold() not in included)
        return order

    def _options(self, options: ResumeRenderOptions | None) -> DocxRenderOptions:
        if options is None:
            return DocxRenderOptions()
        if isinstance(options, DocxRenderOptions):
            return options
        return DocxRenderOptions.model_validate(options.model_dump())

    def _output_path(self, requested: Path, rendered: RenderedResume) -> Path:
        if requested.exists() and requested.is_dir():
            return requested / generate_resume_filename(
                rendered.candidate_name,
                rendered.target_job_title,
                rendered.target_company,
                extension=".docx",
            )
        suffix = requested.suffix.casefold()
        if suffix and suffix != ".docx":
            raise UnsupportedResumeFormatError("DOCX renderer only supports .docx output")
        name = sanitize_filename(requested.name or "resume.docx")
        if not Path(name).suffix:
            name += ".docx"
        return requested.parent / name

    def _next_available_path(self, path: Path) -> Path:
        version = 2
        while True:
            candidate = path.with_name(f"{path.stem}_{version}{path.suffix}")
            if not candidate.exists():
                return candidate
            version += 1

    def _limited_entries[EntryT](
        self,
        entries: Sequence[EntryT],
        limit: int | None,
        label: str,
        warnings: list[str],
    ) -> list[EntryT]:
        if limit is None or len(entries) <= limit:
            return list(entries)
        warnings.append(f"trimmed {len(entries) - limit} {label} entry(s)")
        return list(entries[:limit])

    def _limited_bullets(
        self,
        bullets: Sequence[str],
        limit: int | None,
        label: str,
        warnings: list[str],
    ) -> list[str]:
        if limit is None or len(bullets) <= limit:
            return list(bullets)
        warnings.append(f"trimmed {len(bullets) - limit} bullet(s) from {label}")
        return list(bullets[:limit])

    def _deduplicate_skills(self, skills: Iterable[Skill]) -> list[Skill]:
        result: list[Skill] = []
        seen: set[str] = set()
        for skill in skills:
            canonical = canonical_skill_name(skill.name)
            if canonical not in seen:
                seen.add(canonical)
                result.append(skill)
        return result

    def _dates(self, start: str | None, end: str | None) -> str:
        return " - ".join(value for value in [start, end] if value)

    def _deduplicate(self, values: Iterable[str]) -> list[str]:
        result: list[str] = []
        seen: set[str] = set()
        for value in values:
            if value not in seen:
                seen.add(value)
                result.append(value)
        return result

    def _validate_resume(self, resume: object) -> None:
        if not isinstance(resume, (Resume, OptimizedResume)):
            raise ResumeRenderingError("resume must be a Resume or OptimizedResume")
