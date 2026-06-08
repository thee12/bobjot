"""Plain-text document extraction for supported resume source files.

This layer is intentionally limited to extracting raw text from PDF and DOCX
files. It does not perform AI parsing, infer resume sections, or create Resume
domain objects.
"""

from pathlib import Path

import pdfplumber
from docx import Document

from ai_internship_assistant.domain.models.common import FileFormat


class UnsupportedDocumentFormatError(ValueError):
    """Raised when a file extension is not supported for text extraction."""


class DocumentTextExtractor:
    """Extract plain text from supported document formats."""

    _FORMAT_BY_SUFFIX: dict[str, FileFormat] = {
        ".pdf": FileFormat.PDF,
        ".docx": FileFormat.DOCX,
    }

    def extract_text(self, file_path: str | Path) -> str:
        """Return plain text from a PDF or DOCX document."""

        path = Path(file_path)
        if not path.exists():
            msg = f"document does not exist: {path}"
            raise FileNotFoundError(msg)
        if not path.is_file():
            msg = f"document path is not a file: {path}"
            raise IsADirectoryError(msg)

        file_format = self.detect_format(path)
        if file_format is FileFormat.PDF:
            return self._extract_pdf_text(path)
        if file_format is FileFormat.DOCX:
            return self._extract_docx_text(path)

        msg = f"unsupported document format: {path.suffix}"
        raise UnsupportedDocumentFormatError(msg)

    def detect_format(self, file_path: str | Path) -> FileFormat:
        """Infer the supported document format from the file extension."""

        suffix = Path(file_path).suffix.lower()
        try:
            return self._FORMAT_BY_SUFFIX[suffix]
        except KeyError as exc:
            msg = f"unsupported document format: {suffix or '<none>'}"
            raise UnsupportedDocumentFormatError(msg) from exc

    def _extract_pdf_text(self, path: Path) -> str:
        """Extract page text from a PDF using pdfplumber."""

        page_text: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                cleaned_text = text.strip()
                if cleaned_text:
                    page_text.append(cleaned_text)

        return "\n\n".join(page_text)

    def _extract_docx_text(self, path: Path) -> str:
        """Extract paragraph and table text from a DOCX file using python-docx."""

        document = Document(path)
        blocks: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)

        for table in document.tables:
            for row in table.rows:
                cell_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cell_text:
                    blocks.append("\t".join(cell_text))

        return "\n".join(blocks)


def extract_text(file_path: str | Path) -> str:
    """Convenience function for extracting plain text from a supported document."""

    return DocumentTextExtractor().extract_text(file_path)

