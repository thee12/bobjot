"""Tests for plain-text document extraction from sample files."""

from pathlib import Path

import pytest
from docx import Document

from ai_internship_assistant.services.document_extraction import (
    CorruptedDocumentError,
    DocumentTextExtractor,
    UnsupportedDocumentFormatError,
    extract_text,
)


def _pdf_escape(text: str) -> str:
    """Escape text for a simple PDF string literal."""

    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_sample_pdf(path: Path, lines: list[str]) -> None:
    """Write a small text-only PDF sample for extraction tests."""

    text_commands = []
    for index, line in enumerate(lines):
        y_position = 720 - (index * 18)
        text_commands.append(f"BT /F1 12 Tf 72 {y_position} Td ({_pdf_escape(line)}) Tj ET")

    stream = "\n".join(text_commands).encode("ascii")
    objects = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        (
            b"3 0 obj\n"
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\n"
            b"endobj\n"
        ),
        b"4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
        b"5 0 obj\n<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n"
        + stream
        + b"\nendstream\nendobj\n",
    ]

    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for pdf_object in objects:
        offsets.append(len(content))
        content.extend(pdf_object)

    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )

    path.write_bytes(content)


def _write_sample_docx(path: Path) -> None:
    """Write a small DOCX sample for extraction tests."""

    document = Document()
    document.add_paragraph("Example Candidate")
    document.add_paragraph("Python development internship experience")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "FastAPI"
    document.save(path)


def test_extracts_plain_text_from_sample_pdf(tmp_path: Path) -> None:
    """PDF extraction should return text from a real sample PDF file."""

    sample_pdf = tmp_path / "sample_resume.pdf"
    _write_sample_pdf(sample_pdf, ["Example Candidate", "Python and SQL"])

    text = DocumentTextExtractor().extract_text(sample_pdf)

    assert "Example Candidate" in text
    assert "Python and SQL" in text


def test_extracts_plain_text_from_sample_docx(tmp_path: Path) -> None:
    """DOCX extraction should return paragraph and table text from a sample file."""

    sample_docx = tmp_path / "sample_resume.docx"
    _write_sample_docx(sample_docx)

    text = extract_text(sample_docx)

    assert "Example Candidate" in text
    assert "Python development internship experience" in text
    assert "Skill\tFastAPI" in text


def test_unsupported_document_format_is_rejected(tmp_path: Path) -> None:
    """Unsupported file types should fail before any extraction attempt."""

    text_file = tmp_path / "resume.txt"
    text_file.write_text("plain text is not a supported source file", encoding="utf-8")

    with pytest.raises(UnsupportedDocumentFormatError):
        DocumentTextExtractor().extract_text(text_file)


def test_missing_document_is_rejected(tmp_path: Path) -> None:
    """Missing files should raise the standard file-not-found error."""

    with pytest.raises(FileNotFoundError):
        DocumentTextExtractor().extract_text(tmp_path / "missing.pdf")


def test_corrupted_pdf_is_rejected(tmp_path: Path) -> None:
    """Unreadable PDF files should raise a meaningful extraction error."""

    corrupted_pdf = tmp_path / "corrupted.pdf"
    corrupted_pdf.write_bytes(b"not a real pdf")

    with pytest.raises(CorruptedDocumentError):
        DocumentTextExtractor().extract_text(corrupted_pdf)


def test_invalid_docx_is_rejected(tmp_path: Path) -> None:
    """Unreadable DOCX files should raise a meaningful extraction error."""

    invalid_docx = tmp_path / "invalid.docx"
    invalid_docx.write_bytes(b"not a real docx")

    with pytest.raises(CorruptedDocumentError):
        DocumentTextExtractor().extract_text(invalid_docx)
