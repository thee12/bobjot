"""Tests for deterministic ATS-friendly DOCX resume export."""

from pathlib import Path

import pytest
from docx import Document

from ai_internship_assistant.domain.models import (
    DocxRenderOptions,
    OptimizedResume,
    Resume,
    ResumeOutputFormat,
)
from ai_internship_assistant.services import (
    DocxResumeRenderer,
    ResumeOutputFileExistsError,
    UnsupportedResumeFormatError,
)
from tests.test_resume_renderer import _optimized, _resume


def _document_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _export(
    tmp_path: Path,
    resume: Resume | OptimizedResume | None = None,
    options: DocxRenderOptions | None = None,
) -> tuple[Path, str]:
    result = DocxResumeRenderer().render_to_file(resume or _resume(), tmp_path, options)
    path = Path(result.path)
    return path, _document_text(path)


def test_export_basic_resume_to_docx(tmp_path: Path) -> None:
    path, text = _export(tmp_path)

    assert path.exists()
    assert path.suffix == ".docx"
    assert "Daniel Flores" in text


def test_export_optimized_resume_to_docx(tmp_path: Path) -> None:
    path, text = _export(tmp_path, _optimized())

    assert path.name == "alex_candidate_soc_analyst_intern_example_security_resume.docx"
    assert "Packet Sniffer Project" in text


def test_candidate_contact_and_sections_are_present(tmp_path: Path) -> None:
    _, text = _export(tmp_path)

    assert "daniel@example.com" in text
    assert "linkedin.com/in/danielflores" in text
    assert "EDUCATION" in text
    assert "CERTIFICATIONS" in text
    assert "SKILLS" in text
    assert "PROJECTS" in text
    assert "EXPERIENCE" in text


def test_contact_fields_render_only_when_present(tmp_path: Path) -> None:
    resume = Resume(full_name="Daniel Flores", email="daniel@example.com")

    _, text = _export(tmp_path, resume)

    assert "daniel@example.com" in text
    assert "None" not in text
    assert "LinkedIn" not in text


def test_education_certification_and_skills_render_exactly(tmp_path: Path) -> None:
    _, text = _export(tmp_path)

    assert "B.S., Computer Science - North Carolina State University" in text
    assert "CompTIA Security+ - CompTIA" in text
    assert "Programming: Python, Java" in text
    assert "Security: Linux, Networking" in text


def test_projects_experience_bullets_and_technologies_render(tmp_path: Path) -> None:
    path, text = _export(tmp_path)
    document = Document(path)

    assert "Packet Sniffer Project" in text
    assert "Technologies: Python, Linux" in text
    assert "Built a Python packet sniffer using raw sockets." in text
    assert "IT Assistant - Campus IT" in text
    assert any(
        paragraph.style.name == "List Bullet"
        and paragraph.text == "Built a Python packet sniffer using raw sockets."
        for paragraph in document.paragraphs
    )


def test_additional_sections_render_generically(tmp_path: Path) -> None:
    optimized = _optimized().model_copy(
        update={"additional_sections": {"Leadership": ["Led student security club."]}}
    )

    _, text = _export(tmp_path, optimized)

    assert "LEADERSHIP" in text
    assert "Led student security club." in text


def test_missing_sections_are_omitted_gracefully(tmp_path: Path) -> None:
    resume = Resume(full_name="Daniel Flores", email="daniel@example.com")

    _, text = _export(tmp_path, resume)

    assert "EDUCATION" not in text
    assert "PROJECTS" not in text


def test_exporter_does_not_add_unsafe_content_or_rewrite_bullets(tmp_path: Path) -> None:
    _, text = _export(tmp_path)

    assert "Built a Python packet sniffer using raw sockets." in text
    assert "Splunk" not in text
    assert "SIEM" not in text
    assert "40%" not in text


def test_max_bullets_and_entries_are_respected(tmp_path: Path) -> None:
    options = DocxRenderOptions(
        max_projects=1,
        max_experiences=0,
        max_bullets_per_project=1,
    )
    result = DocxResumeRenderer().render_to_file(_resume(), tmp_path, options)
    text = _document_text(Path(result.path))

    assert "Packet Sniffer Project" in text
    assert "Portfolio" not in text
    assert "Documented packet fields." not in text
    assert "EXPERIENCE" not in text
    assert any("trimmed" in warning for warning in result.rendered_resume.warnings)


def test_summary_visibility_is_configurable(tmp_path: Path) -> None:
    _, included = _export(tmp_path / "included", options=DocxRenderOptions(include_summary=True))
    _, omitted = _export(tmp_path / "omitted", options=DocxRenderOptions(include_summary=False))

    assert "SUMMARY" in included
    assert "SUMMARY" not in omitted


def test_parent_directory_is_created_and_filename_is_sanitized(tmp_path: Path) -> None:
    output = tmp_path / "nested" / "Daniel: Resume?.DOCX"

    result = DocxResumeRenderer().render_to_file(_resume(), output)

    assert Path(result.path).exists()
    assert Path(result.path).name == "daniel_resume.docx"


def test_overwrite_protection_and_allow_overwrite(tmp_path: Path) -> None:
    output = tmp_path / "resume.docx"
    renderer = DocxResumeRenderer()
    renderer.render_to_file(_resume(), output)

    with pytest.raises(ResumeOutputFileExistsError):
        renderer.render_to_file(_resume(), output)

    result = renderer.render_to_file(
        _resume(),
        output,
        DocxRenderOptions(allow_overwrite=True),
    )
    assert Path(result.path).exists()


def test_directory_output_avoids_duplicate_names(tmp_path: Path) -> None:
    renderer = DocxResumeRenderer()
    first = renderer.render_to_file(_resume(), tmp_path)
    second = renderer.render_to_file(_resume(), tmp_path)

    assert Path(first.path).name == "daniel_flores_resume_resume.docx"
    assert Path(second.path).name == "daniel_flores_resume_resume_2.docx"


def test_rendered_file_metadata_is_returned(tmp_path: Path) -> None:
    result = DocxResumeRenderer().render_to_file(_resume(), tmp_path)

    assert result.format == ResumeOutputFormat.DOCX
    assert result.byte_size > 0
    assert len(result.content_hash) == 64
    assert result.rendered_resume.format == ResumeOutputFormat.DOCX
    assert result.rendered_resume.renderer_version == "docx-resume-renderer-v1"


def test_docx_uses_no_tables_images_headers_or_footers(tmp_path: Path) -> None:
    path, _ = _export(tmp_path)
    document = Document(path)

    assert document.tables == []
    assert len(document.inline_shapes) == 0
    assert all(not section.header.paragraphs[0].text for section in document.sections)
    assert all(not section.footer.paragraphs[0].text for section in document.sections)


def test_docx_style_defaults_are_ats_friendly(tmp_path: Path) -> None:
    path, _ = _export(tmp_path)
    document = Document(path)
    normal = document.styles["Normal"]

    assert normal.font.name == "Calibri"
    assert normal.font.size.pt == 10.5
    assert document.sections[0].left_margin.inches == pytest.approx(0.6, abs=0.01)
    assert document.sections[0].right_margin.inches == pytest.approx(0.6, abs=0.01)


def test_custom_typography_and_margins_are_applied(tmp_path: Path) -> None:
    options = DocxRenderOptions(
        font_name="Arial",
        font_size=11,
        margin_left=0.7,
        margin_right=0.7,
    )

    path, _ = _export(tmp_path, options=options)
    document = Document(path)

    assert document.styles["Normal"].font.name == "Arial"
    assert document.styles["Normal"].font.size.pt == 11
    assert document.sections[0].left_margin.inches == pytest.approx(0.7, abs=0.01)


def test_unsupported_extension_is_rejected(tmp_path: Path) -> None:
    with pytest.raises(UnsupportedResumeFormatError):
        DocxResumeRenderer().render_to_file(_resume(), tmp_path / "resume.pdf")


def test_original_resume_is_not_mutated(tmp_path: Path) -> None:
    resume = _resume()
    before = resume.model_dump()

    DocxResumeRenderer().render_to_file(
        resume,
        tmp_path,
        DocxRenderOptions(max_projects=1, max_bullets_per_project=1),
    )

    assert resume.model_dump() == before


def test_render_returns_visible_text_without_writing_file() -> None:
    rendered = DocxResumeRenderer().render(_resume())

    assert rendered.format == ResumeOutputFormat.DOCX
    assert "Daniel Flores" in rendered.content
    assert "Packet Sniffer Project" in rendered.content


def test_source_ids_are_preserved_in_metadata_not_document(tmp_path: Path) -> None:
    options = DocxRenderOptions(source_resume_id="master-1", source_version_id="version-1")

    result = DocxResumeRenderer().render_to_file(_resume(), tmp_path, options)
    text = _document_text(Path(result.path))

    assert result.rendered_resume.source_resume_id == "master-1"
    assert result.rendered_resume.source_version_id == "version-1"
    assert "master-1" not in text
    assert "version-1" not in text


def test_custom_section_order_is_respected(tmp_path: Path) -> None:
    options = DocxRenderOptions(section_order=["Experience", "Education"])

    _, text = _export(tmp_path, options=options)

    assert text.index("EXPERIENCE") < text.index("EDUCATION")


def test_missing_name_and_empty_sections_produce_warnings() -> None:
    rendered = DocxResumeRenderer().render(Resume(email="daniel@example.com"))

    assert "candidate name is missing" in rendered.warnings
    assert "resume contains no renderable sections" in rendered.warnings


def test_same_input_produces_same_visible_text() -> None:
    renderer = DocxResumeRenderer()

    first = renderer.render(_resume())
    second = renderer.render(_resume())

    assert first.content == second.content
    assert first.warnings == second.warnings
